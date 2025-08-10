import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import plotly.express as px
from io import BytesIO
import chardet
import pytz
from bs4 import BeautifulSoup
import base64
import textwrap
from datetime import datetime
import os

# ====== (opcional) DOCX ======
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_OK = True
except Exception:
    DOCX_OK = False

# ====== (opcional) PDF ======
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.lib.utils import ImageReader
    PDF_OK = True
except Exception:
    PDF_OK = False

st.set_page_config(page_title="Dashboard Inteligente", layout="wide")
st.title("Dashboard Inteligente - HTML, XLSX e CSV")

# =============================
# Funções auxiliares
# =============================

def _decode_file(uploaded_file) -> str:
    raw = uploaded_file.read()
    enc = (chardet.detect(raw)["encoding"] or "utf-8")
    try:
        text = raw.decode(enc, errors="replace")
    except Exception:
        text = raw.decode("utf-8", errors="replace")
    finally:
        uploaded_file.seek(0)
    return text

def _parse_text_time_ip(text: str) -> pd.DataFrame | None:
    soup = BeautifulSoup(text, "html.parser")
    plain = soup.get_text(separator="\n")
    lines = [ln.strip() for ln in plain.splitlines()]
    skip_prefixes = ("WhatsApp Business Record Page",)
    clean = [ln for ln in lines if ln and not any(ln.startswith(p) for p in skip_prefixes)]

    records = []
    current = {"Time": None, "IP Address": None}

    i, n = 0, len(clean)
    while i < n:
        token = clean[i]
        if token.lower() == "time":
            j = i + 1
            while j < n and not clean[j]:
                j += 1
            if j < n:
                current["Time"] = clean[j]
                i = j
        elif token.lower() in ("ip address", "ip addresses"):
            j = i + 1
            while j < n and not clean[j]:
                j += 1
            if j < n:
                current["IP Address"] = clean[j]
                i = j
        else:
            if token.lower() in ("ip addresses", "ipaddress", "ip"):
                j = i + 1
                while j < n and not clean[j]:
                    j += 1
                if j < n:
                    current["IP Address"] = clean[j]
                    i = j

        if current["Time"] and current["IP Address"]:
            records.append({"Time": current["Time"], "IP Address": current["IP Address"]})
            current = {"Time": None, "IP Address": None}
        i += 1

    if not records:
        return None

    df = pd.DataFrame(records, columns=["Time", "IP Address"])
    try:
        df["Time"] = pd.to_datetime(df["Time"], errors="coerce", utc=True)
    except Exception:
        pass
    return df

def ler_arquivo(uploaded_file):
    ext = uploaded_file.name.split('.')[-1].lower()

    if ext == "xlsx":
        try:
            uploaded_file.seek(0)
            return pd.read_excel(uploaded_file)
        finally:
            uploaded_file.seek(0)

    elif ext == "csv":
        try:
            raw = uploaded_file.read()
            enc = (chardet.detect(raw)["encoding"] or "utf-8")
            st.info(f"Arquivo CSV detectado com encoding: **{enc}**")
            uploaded_file.seek(0)
            return pd.read_csv(uploaded_file, sep=None, engine="python", encoding=enc)
        except Exception as e:
            st.error(f"Erro ao ler CSV: {e}")
            return None
        finally:
            uploaded_file.seek(0)

    elif ext in ("html", "htm", "txt"):
        try:
            text = _decode_file(uploaded_file)

            soup = BeautifulSoup(text, "html.parser")
            if soup.find("table"):
                try:
                    tables = pd.read_html(text, flavor="bs4")
                    if tables:
                        return tables[0]
                except Exception:
                    pass

            df_text = _parse_text_time_ip(text)
            if df_text is not None and not df_text.empty:
                return df_text

            st.error("Não foi possível extrair dados: sem <table> e formato não corresponde a Time/IP.")
            return None
        except Exception as e:
            st.error(f"Erro ao ler arquivo de texto/HTML: {e}")
            return None
        finally:
            uploaded_file.seek(0)

    else:
        st.warning("Extensão não suportada. Use HTML/HTM/TXT, XLSX ou CSV.")
        return None

def to_excel(df):
    df_copy = df.copy()
    for col in df_copy.columns:
        if pd.api.types.is_datetime64_any_dtype(df_copy[col]):
            if getattr(df_copy[col].dt, "tz", None) is not None:
                df_copy[col] = df_copy[col].dt.tz_localize(None)
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_copy.to_excel(writer, index=False)
    return output.getvalue()

def to_csv(df):
    return df.to_csv(index=False).encode('utf-8')

def to_json(df):
    return df.to_json(orient="records", force_ascii=False).encode('utf-8')

def gerar_insights(df):
    insights = []
    insights.append(f"O conjunto de dados possui {df.shape[0]} linhas e {df.shape[1]} colunas.")
    insights.append(f"As colunas disponíveis são: {', '.join(map(str, df.columns))}.")
    num_cols = df.select_dtypes(include='number').columns
    if len(num_cols) > 0:
        insights.append(f"Foram encontradas {len(num_cols)} colunas numéricas.")
        for col in num_cols:
            try:
                media = df[col].mean()
                insights.append(f"A média da coluna '{col}' é {media:.2f}.")
            except Exception:
                pass
    else:
        insights.append("Não há colunas numéricas para calcular estatísticas básicas.")
    cat_cols = df.select_dtypes(exclude='number').columns
    if len(cat_cols) > 0:
        insights.append(f"Foram encontradas {len(cat_cols)} colunas categóricas.")
        for col in cat_cols:
            try:
                modo = df[col].mode()
                valor_mais_freq = modo.iloc[0] if not modo.empty else "Nenhum"
                insights.append(f"Na coluna '{col}', o valor mais frequente é '{valor_mais_freq}'.")
            except Exception:
                pass
    return "\n".join(insights)

def converter_datas_para_timestamp(df):
    df_convertido = df.copy()
    for col in df_convertido.columns:
        if pd.api.types.is_datetime64_any_dtype(df_convertido[col]):
            df_convertido[col] = df_convertido[col].view('int64') / 1e9
    return df_convertido

def detectar_colunas_datetime(df):
    fuso = pytz.timezone("America/Sao_Paulo")
    df = df.copy()
    for col in df.columns:
        if df[col].dtype == object:
            try:
                temp = pd.to_datetime(df[col], errors='raise', utc=True)
                temp = temp.dt.tz_convert(fuso)
                df[col] = temp
            except Exception:
                pass
        elif pd.api.types.is_datetime64_any_dtype(df[col]):
            try:
                if getattr(df[col].dt, "tz", None) is not None:
                    df[col] = df[col].dt.tz_convert(fuso)
                else:
                    df[col] = df[col].dt.tz_localize("UTC").dt.tz_convert(fuso)
            except Exception:
                pass
    return df

def formatar_datas_para_exibicao(df):
    df_exibir = df.copy()
    for col in df_exibir.columns:
        if pd.api.types.is_datetime64_any_dtype(df_exibir[col]):
            df_exibir[col] = df_exibir[col].dt.strftime("%d/%m/%Y %H:%M:%S")
    return df_exibir

# ====== Relatório: detecção de colunas e tabelas ======

def _guess_colunas(df):
    col_tempo = None
    for c in df.columns:
        if str(c).strip().lower() == "time":
            col_tempo = c
            break
    if col_tempo is None:
        for c in df.columns:
            if pd.api.types.is_datetime64_any_dtype(df[c]):
                col_tempo = c
                break
    if col_tempo is None:
        candidatos_tempo = [c for c in df.columns if str(c).strip().lower() in
                            ("timestamp","data","datetime","date","hora","data/hora")]
        for cand in candidatos_tempo:
            try:
                test = pd.to_datetime(df[cand], errors="coerce", utc=True)
                if test.notna().any():
                    col_tempo = cand
                    break
            except Exception:
                pass

    col_ip = None
    for c in df.columns:
        if str(c).strip().lower() == "ip address":
            col_ip = c
            break
    if col_ip is None:
        candidatos_ip = [c for c in df.columns if "ip" in str(c).lower()]
        if candidatos_ip:
            col_ip = candidatos_ip[0]

    return col_tempo, col_ip

def _fig_to_png_bytes(fig):
    buf = BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=180)
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()

def _grafico_timeline(df, col_tempo):
    serie = pd.to_datetime(df[col_tempo], errors="coerce")
    serie = serie.dropna()
    if serie.empty:
        return None
    por_dia = serie.dt.date.value_counts().sort_index()
    fig, ax = plt.subplots()
    ax.plot(por_dia.index, por_dia.values, marker="o")
    ax.set_title("Linha do tempo de eventos por dia")
    ax.set_xlabel("Data")
    ax.set_ylabel("Quantidade de eventos")
    ax.grid(True, linewidth=0.3)
    return _fig_to_png_bytes(fig)

def _grafico_top_ips(df, col_ip, top_n=10):
    if col_ip is None or col_ip not in df.columns:
        return None
    cont = df[col_ip].astype(str).value_counts().head(top_n)
    if cont.empty:
        return None
    fig, ax = plt.subplots()
    cont.plot(kind="barh", ax=ax)
    ax.invert_yaxis()
    ax.set_title(f"Top {min(top_n, len(cont))} IPs por frequência")
    ax.set_xlabel("Ocorrências")
    return _fig_to_png_bytes(fig)

def _png_data_uri(png_bytes):
    b64 = base64.b64encode(png_bytes).decode("ascii")
    return f"data:image/png;base64,{b64}"

def montar_tabela_ip_time_completa(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame(columns=["Time (America/Sao_Paulo)", "IP Address"])

    col_tempo, col_ip = _guess_colunas(df)
    if not col_tempo or not col_ip or col_tempo not in df.columns or col_ip not in df.columns:
        return pd.DataFrame(columns=["Time (America/Sao_Paulo)", "IP Address"])

    base = df[[col_tempo, col_ip]].copy()

    try:
        serie = pd.to_datetime(base[col_tempo], errors="coerce", utc=True)
        serie = serie.dt.tz_convert("America/Sao_Paulo")
    except Exception:
        serie = pd.to_datetime(base[col_tempo], errors="coerce")
        if getattr(serie.dt, "tz", None) is None:
            serie = serie.dt.tz_localize("America/Sao_Paulo")

    base["__time"] = serie
    base = base.dropna(subset=["__time", col_ip])

    base = base.sort_values("__time", ascending=False)

    base["Time (America/Sao_Paulo)"] = base["__time"].dt.strftime("%d/%m/%Y %H:%M:%S")
    base = base[["Time (America/Sao_Paulo)", col_ip]]
    base = base.rename(columns={col_ip: "IP Address"})
    return base.reset_index(drop=True)

def gerar_relatorio_html_txt_docx(df_base: pd.DataFrame,
                                  df_filtrado: pd.DataFrame,
                                  incluir_graficos: bool,
                                  metadados: dict):
    df = df_filtrado if df_filtrado is not None and not df_filtrado.empty else df_base.copy()

    col_tempo, col_ip = _guess_colunas(df)

    periodo_txt = "Não identificado"
    if col_tempo and col_tempo in df.columns:
        try:
            serie = pd.to_datetime(df[col_tempo], errors="coerce", utc=True)
            if getattr(serie.dt, "tz", None) is not None:
                serie = serie.dt.tz_convert("America/Sao_Paulo")
            else:
                serie = serie.dt.tz_localize("UTC").dt.tz_convert("America/Sao_Paulo")
            tmin, tmax = serie.min(), serie.max()
            if pd.notna(tmin) and pd.notna(tmax):
                periodo_txt = f"{tmin.strftime('%d/%m/%Y %H:%M:%S')} a {tmax.strftime('%d/%m/%Y %H:%M:%S')}"
        except Exception:
            pass

    achados = []
    achados.append(f"Total de registros analisados: {len(df)}.")
    achados.append(f"Total de colunas: {df.shape[1]} ({', '.join(map(str, df.columns))}).")
    achados.append(f"Período coberto (se aplicável): {periodo_txt}.")
    if col_ip and col_ip in df.columns:
        ips_unicos = df[col_ip].astype(str).nunique(dropna=True)
        achados.append(f"Endereços IP distintos identificados: {ips_unicos}.")
        top_ips = df[col_ip].astype(str).value_counts().head(5)
        if not top_ips.empty:
            resumo_top = "; ".join([f"{idx} ({val})" for idx, val in top_ips.items()])
            achados.append(f"Principais IPs por frequência: {resumo_top}.")
    else:
        achados.append("Não foi identificada coluna de IP.")

    png_timeline = _grafico_timeline(df, col_tempo) if incluir_graficos and col_tempo else None
    png_top_ips = _grafico_top_ips(df, col_ip) if incluir_graficos and col_ip else None

    tabela_completa = montar_tabela_ip_time_completa(df)

    html_parts = []
    html_parts.append("<meta charset='utf-8'>")
    html_parts.append("<style>body{font-family:Arial,Helvetica,sans-serif;margin:24px} h1,h2{margin:0.2em 0} table{border-collapse:collapse;width:100%} th,td{border:1px solid #ddd;padding:6px;font-size:13px} .muted{color:#555} .blk{margin:18px 0}</style>")
    html_parts.append("<h1>Relatório Policial - Análise de IPs (Análise de Dados)</h1>")
    html_parts.append("<div class='blk'><h2>Metadados</h2><table>")
    for k,v in metadados.items():
        html_parts.append(f"<tr><th style='width:260px;text-align:left'>{k}</th><td>{v}</td></tr>")
    html_parts.append("</table></div>")
    html_parts.append("<div class='blk'><h2>Síntese dos Achados</h2><ul>")
    for a in achados:
        html_parts.append(f"<li>{a}</li>")
    html_parts.append("</ul></div>")
    html_parts.append("<div class='blk'><h2>Metodologia</h2>")
    html_parts.append("<p class='muted'>Os dados foram importados, higienizados e analisados com apoio de ferramentas computacionais. Procedeu-se à consolidação de múltiplas fontes, conversão de datas para o fuso America/Sao_Paulo e análise descritiva (contagens, modos e médias).</p></div>")

    if incluir_graficos and (png_timeline or png_top_ips):
        html_parts.append("<div class='blk'><h2>Gráficos</h2>")
        if png_timeline:
            html_parts.append("<h3>Linha do tempo de eventos por dia</h3>")
            html_parts.append(f"<img src='{_png_data_uri(png_timeline)}' style='max-width:100%;height:auto'/>")
        if png_top_ips:
            html_parts.append("<h3>Top IPs por frequência</h3>")
            html_parts.append(f"<img src='{_png_data_uri(png_top_ips)}' style='max-width:100%;height:auto'/>")
        html_parts.append("</div>")

    html_parts.append("<div class='blk'><h2>Tabela Completa: IP Address × Time (mais recentes primeiro)</h2>")
    if tabela_completa.empty:
        html_parts.append("<p class='muted'>Não há dados suficientes para compor a tabela completa (verifique colunas de IP e horário).</p>")
    else:
        html_parts.append(tabela_completa.to_html(index=False))
    html_parts.append("</div>")

    html_bytes = "\n".join(html_parts).encode("utf-8")

    linhas = []
    linhas.append("RELATÓRIO POLICIAL (ANÁLISE DE DADOS)")
    linhas.append("=" * 60)
    for k,v in metadados.items():
        linhas.append(f"{k}: {v}")
    linhas.append("")
    linhas.append("1. SÍNTESE DOS ACHADOS")
    for linha in achados:
        linhas.append(f"- {linha}")
    linhas.append("")
    linhas.append("2. METODOLOGIA")
    linhas.append(textwrap.fill(
        "Os dados foram importados, higienizados e analisados com apoio de ferramentas "
        "computacionais. Procedeu-se à consolidação de múltiplas fontes, conversão de datas "
        "para o fuso America/Sao_Paulo e análise descritiva (contagens, modos e médias).",
        width=100
    ))
    txt_bytes = "\n".join(linhas).encode("utf-8")

    docx_bytes = None
    if DOCX_OK:
        doc = Document()
        doc.add_heading('Relatório Policial - Análise de IPs (Análise de Dados)', level=1)

        doc.add_heading('Metadados', level=2)
        for k,v in metadados.items():
            doc.add_paragraph(f"{k}: {v}")

        doc.add_heading('Síntese dos Achados', level=2)
        for a in achados:
            doc.add_paragraph(a)

        doc.add_heading('Metodologia', level=2)
        doc.add_paragraph(
            "Os dados foram importados, higienizados e analisados com apoio de ferramentas computacionais. "
            "Procedeu-se à consolidação de múltiplas fontes, conversão de datas para o fuso America/Sao_Paulo "
            "e análise descritiva (contagens, modos e médias)."
        )

        if incluir_graficos and (png_timeline or png_top_ips):
            doc.add_heading('Gráficos', level=2)
            if png_timeline:
                doc.add_paragraph("Linha do tempo de eventos por dia")
                stream = BytesIO(png_timeline); stream.seek(0)
                doc.add_picture(stream, width=Inches(6.0))
            if png_top_ips:
                doc.add_paragraph("Top IPs por frequência")
                stream = BytesIO(png_top_ips); stream.seek(0)
                doc.add_picture(stream, width=Inches(6.0))

        doc.add_heading('Tabela Completa: IP Address × Time (mais recentes primeiro)', level=2)
        tabela = montar_tabela_ip_time_completa(df)
        if tabela.empty:
            doc.add_paragraph("Não há dados suficientes para compor a tabela completa (verifique colunas de IP e horário).")
        else:
            cols = ["Time (America/Sao_Paulo)", "IP Address"]
            t = doc.add_table(rows=1, cols=len(cols))
            hdr = t.rows[0].cells
            for i, c in enumerate(cols):
                hdr[i].text = c
            for _, row in tabela.iterrows():
                cells = t.add_row().cells
                cells[0].text = str(row["Time (America/Sao_Paulo)"])
                cells[1].text = str(row["IP Address"])

        bio = BytesIO()
        doc.save(bio); bio.seek(0)
        docx_bytes = bio.getvalue()

    return {"html": html_bytes, "txt": txt_bytes, "docx": docx_bytes}

# ---- PDF: brasão + numeração + gráficos; título 2 cm abaixo do brasão (via topMargin) ----
def _header_footer(canvas, doc):
    largura_pagina, altura_pagina = A4
    try:
        logo_path = "brasao.png"
        if os.path.exists(logo_path):
            img_w = 70 * mm      # largura do brasão
            img_h = 70 * mm      # altura do brasão
            x = (largura_pagina - img_w) / 2.5
            y = altura_pagina - (img_h + 10 * mm)  # topo do brasão a 10 mm do topo da página
            canvas.drawImage(logo_path, x, y, width=img_w, height=img_h,
                             preserveAspectRatio=True, mask='auto')
    except Exception:
        pass
    # Rodapé com numeração
    page_num = canvas.getPageNumber()
    canvas.setFont("Helvetica", 9)
    canvas.setFillColor(colors.grey)
    canvas.drawRightString(largura_pagina - 20 * mm, 12 * mm, f"Página {page_num}")

def _rl_image_from_png_bytes(png_bytes: bytes, max_width_pt: float, max_height_pt: float):
    try:
        bio = BytesIO(png_bytes)
        ir = ImageReader(bio)
        iw, ih = ir.getSize()
        scale = min(max_width_pt / float(iw), max_height_pt / float(ih), 1.0)
        w = iw * scale
        h = ih * scale
        bio.seek(0)
        return Image(bio, width=w, height=h)
    except Exception:
        return None

def gerar_relatorio_pdf(df_base: pd.DataFrame,
                        df_filtrado: pd.DataFrame,
                        incluir_graficos: bool,
                        metadados: dict,
                        titulo: str = "Relatório Policial - Análise de IPs") -> bytes:
    if not PDF_OK:
        raise RuntimeError("Pacote 'reportlab' não está disponível. Instale com: pip install reportlab")

    df = df_filtrado if df_filtrado is not None and not df_filtrado.empty else df_base.copy()
    col_tempo, col_ip = _guess_colunas(df)

    periodo_txt = "Não identificado"
    if col_tempo and col_tempo in df.columns:
        try:
            serie = pd.to_datetime(df[col_tempo], errors="coerce", utc=True)
            if getattr(serie.dt, "tz", None) is not None:
                serie = serie.dt.tz_convert("America/Sao_Paulo")
            else:
                serie = serie.dt.tz_localize("UTC").dt.tz_convert("America/Sao_Paulo")
            tmin, tmax = serie.min(), serie.max()
            if pd.notna(tmin) and pd.notna(tmax):
                periodo_txt = f"{tmin.strftime('%d/%m/%Y %H:%M:%S')} a {tmax.strftime('%d/%m/%Y %H:%M:%S')}"
        except Exception:
            pass

    achados = [
        f"Total de registros analisados: {len(df)}.",
        f"Total de colunas: {df.shape[1]} ({', '.join(map(str, df.columns))}).",
        f"Período coberto (se aplicável): {periodo_txt}."
    ]
    if col_ip and col_ip in df.columns:
        ips_unicos = df[col_ip].astype(str).nunique(dropna=True)
        achados.append(f"Endereços IP distintos identificados: {ips_unicos}.")
        top_ips = df[col_ip].astype(str).value_counts().head(5)
        if not top_ips.empty:
            resumo_top = "; ".join([f"{idx} ({val})" for idx, val in top_ips.items()])
            achados.append(f"Principais IPs por frequência: {resumo_top}.")
    else:
        achados.append("Não foi identificada coluna de IP.")

    styles = getSampleStyleSheet()
    style_title = styles["Title"]
    style_h1 = styles["Heading1"]
    style_body = styles["BodyText"]

    # Margens: top = 70mm para reservar 10mm topo + 40mm brasão + 20mm de distância até o título
    left = right = bottom = 36  # ~12,7 mm
    top = 70 * mm               # 70 mm
    frame_width = A4[0] - (left + right)
    frame_height = A4[1] - (top + bottom)

    # Conteúdo (sem Spacer)
    story = []
    story.append(Paragraph(titulo, style_title))

    story.append(Paragraph("<b>Metadados</b>", style_h1))
    for k, v in metadados.items():
        story.append(Paragraph(f"{k}: {v}", style_body))
    story.append(Spacer(1, 8))

    story.append(Paragraph("1. Síntese dos Achados", style_h1))
    for a in achados:
        story.append(Paragraph(a, style_body))
    story.append(Spacer(1, 8))

    story.append(Paragraph("2. Metodologia", style_h1))
    story.append(Paragraph(
        "Os dados foram importados, higienizados e analisados com apoio de ferramentas computacionais. "
        "Procedeu-se à consolidação de múltiplas fontes, conversão de datas para o fuso America/Sao_Paulo "
        "e análise descritiva (contagens, modos e médias).", style_body
    ))
    story.append(Spacer(1, 8))

    if incluir_graficos:
        png_timeline = _grafico_timeline(df, col_tempo) if col_tempo else None
        png_top_ips = _grafico_top_ips(df, col_ip) if col_ip else None

        if png_timeline or png_top_ips:
            story.append(Paragraph("3. Gráficos", style_h1))

            max_w = frame_width
            max_h = frame_height * 0.45

            if png_timeline:
                story.append(Paragraph("Linha do tempo de eventos por dia", style_body))
                img_flow = _rl_image_from_png_bytes(png_timeline, max_w, max_h)
                if img_flow:
                    story.append(Spacer(1, 4))
                    story.append(img_flow)
                    story.append(Spacer(1, 10))

            if png_top_ips:
                story.append(Paragraph("Top IPs por frequência", style_body))
                img_flow = _rl_image_from_png_bytes(png_top_ips, max_w, max_h)
                if img_flow:
                    story.append(Spacer(1, 4))
                    story.append(img_flow)
                    story.append(Spacer(1, 10))

    story.append(Paragraph("4. Tabela Completa: IP Address × Time (mais recentes primeiro)", style_h1))
    tabela_full = montar_tabela_ip_time_completa(df)
    if tabela_full.empty:
        story.append(Paragraph("Não há dados suficientes para compor a tabela completa (verifique colunas de IP e horário).", style_body))
    else:
        data = [list(tabela_full.columns)] + tabela_full.astype(str).values.tolist()
        tbl = Table(data, colWidths=[150, 350])
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
            ("TEXTCOLOR", (0,0), (-1,0), colors.black),
            ("ALIGN", (0,0), (-1,-1), "LEFT"),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("FONTNAME", (0,0), (-1,-1), "Helvetica"),
            ("FONTSIZE", (0,0), (-1,-1), 9),
            ("GRID", (0,0), (-1,-1), 0.25, colors.grey),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.whitesmoke, colors.white]),
        ]))
        story.append(tbl)

    bio = BytesIO()
    doc = SimpleDocTemplate(
        bio,
        pagesize=A4,
        leftMargin=left, rightMargin=right,
        topMargin=top, bottomMargin=bottom,
        title=titulo
    )
    doc.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)
    return bio.getvalue()

# =============================
# Upload múltiplo
# =============================
uploaded_files = st.file_uploader(
    "Selecione arquivos HTML/HTM/TXT, XLSX ou CSV (múltiplos arquivos permitidos)",
    type=["html", "htm", "txt", "xlsx", "csv"],
    accept_multiple_files=True
)

if uploaded_files:
    dfs = []
    for file in uploaded_files:
        df_temp = ler_arquivo(file)
        if df_temp is not None:
            dfs.append(df_temp)

    if dfs:
        df = pd.concat(dfs, ignore_index=True)

        aba1, aba2, aba3, aba4, aba5, aba6, aba7, aba8 = st.tabs([
            "📄 Dados",
            "🔍 Filtros",
            "📊 Gráficos",
            "📈 Estatísticas",
            "📊 Dashboard Automático",
            "🤖 Insights Automáticos",
            "⬇️ Exportações",
            "📝 Relatório Policial"
        ])

        with aba1:
            st.subheader("Visualização dos Dados Combinados")
            df = detectar_colunas_datetime(df)
            st.dataframe(formatar_datas_para_exibicao(df))

        with aba2:
            st.subheader("Filtrar Dados")
            colunas = st.multiselect("Selecione colunas para filtrar", df.columns)
            df_filtrado = df.copy()
            for col in colunas:
                valores = df[col].dropna().unique().tolist()
                selecao = st.multiselect(f"Valores para {col}", valores)
                if selecao:
                    df_filtrado = df_filtrado[df_filtrado[col].isin(selecao)]
            st.dataframe(formatar_datas_para_exibicao(df_filtrado))

        with aba3:
            st.subheader("Visualização de Gráficos")
            if not df_filtrado.empty:
                colunas_num = df_filtrado.select_dtypes(include="number").columns
                colunas_cat = df_filtrado.select_dtypes(exclude="number").columns
                tipo_grafico = st.selectbox(
                    "Selecione o tipo de gráfico",
                    ["Histograma", "Barras", "Linha", "Pizza"]
                )
                modo_grafico = st.radio(
                    "Selecione a biblioteca para visualização",
                    ["Plotly (Interativo)", "Matplotlib (Estático)"]
                )
                if tipo_grafico in ["Histograma", "Linha"] and len(colunas_num) > 0:
                    colunas_escolhidas = st.multiselect("Selecione colunas numéricas", colunas_num)
                elif tipo_grafico in ["Barras", "Pizza"] and len(colunas_cat) > 0:
                    colunas_escolhidas = st.multiselect("Selecione colunas categóricas", colunas_cat)
                else:
                    colunas_escolhidas = []
                if colunas_escolhidas:
                    if modo_grafico == "Plotly (Interativo)":
                        if tipo_grafico == "Histograma":
                            for col in colunas_escolhidas:
                                fig = px.histogram(df_filtrado, x=col, nbins=10, title=f"Histograma - {col}")
                                st.plotly_chart(fig, use_container_width=True)
                        elif tipo_grafico == "Barras":
                            for col in colunas_escolhidas:
                                contagem = df_filtrado[col].value_counts().reset_index()
                                contagem.columns = [col, "Contagem"]
                                fig = px.bar(contagem, x=col, y="Contagem", title=f"Barras - {col}")
                                st.plotly_chart(fig, use_container_width=True)
                        elif tipo_grafico == "Linha":
                            fig = px.line(df_filtrado[colunas_escolhidas])
                            fig.update_layout(title="Gráfico de Linha (múltiplas colunas)")
                            st.plotly_chart(fig, use_container_width=True)
                        elif tipo_grafico == "Pizza":
                            for col in colunas_escolhidas:
                                contagem = df_filtrado[col].value_counts().reset_index()
                                contagem.columns = [col, "Contagem"]
                                fig = px.pie(contagem, names=col, values="Contagem", title=f"Pizza - {col}")
                                st.plotly_chart(fig, use_container_width=True)
                    else:
                        for col in colunas_escolhidas:
                            fig, ax = plt.subplots()
                            if tipo_grafico == "Histograma":
                                df_filtrado[col].plot(kind="hist", bins=10, rwidth=0.8, ax=ax)
                                ax.set_title(f"Histograma - {col}")
                            elif tipo_grafico == "Barras":
                                df_filtrado[col].value_counts().plot(kind="bar", ax=ax)
                                ax.set_title(f"Barras - {col}")
                            elif tipo_grafico == "Linha":
                                df_filtrado[col].plot(kind="line", ax=ax)
                                ax.set_title(f"Linha - {col}")
                            elif tipo_grafico == "Pizza":
                                df_filtrado[col].value_counts().plot(kind="pie", autopct='%1.1f%%', ax=ax)
                                ax.set_ylabel('')
                                ax.set_title(f"Pizza - {col}")
                            st.pyplot(fig)
                else:
                    st.info("Selecione pelo menos uma coluna para gerar o gráfico.")

        with aba4:
            st.subheader("Estatísticas Descritivas")
            if not df_filtrado.empty:
                df_filtrado = detectar_colunas_datetime(df_filtrado)
                st.write("**Estatísticas das colunas numéricas:**")
                st.dataframe(df_filtrado.describe())
                st.write("**Contagem de valores por coluna:**")
                st.dataframe(df_filtrado.count())
                usar_timestamp = st.checkbox(
                    "Converter colunas de datas (datetime) em valores numéricos (timestamp) para incluir na correlação"
                )
                df_corr = df_filtrado.copy()
                if usar_timestamp:
                    df_corr = converter_datas_para_timestamp(df_corr)
                colunas_numericas = df_corr.select_dtypes(include="number")
                if colunas_numericas.shape[1] > 1:
                    st.write("**Matriz de Correlação:**")
                    st.dataframe(colunas_numericas.corr())
                else:
                    st.info("Não há colunas numéricas suficientes para calcular a correlação.")
            else:
                st.info("Nenhum dado disponível para gerar estatísticas.")

        with aba5:
            st.subheader("Dashboard Automático")
            if not df_filtrado.empty:
                colunas_num = df_filtrado.select_dtypes(include="number").columns
                colunas_cat = df_filtrado.select_dtypes(exclude="number").columns
                if len(colunas_num) > 0:
                    col = colunas_num[0]
                    st.write(f"Histograma automático para {col}")
                    st.plotly_chart(px.histogram(df_filtrado, x=col), use_container_width=True)
                if len(colunas_cat) > 0:
                    col = colunas_cat[0]
                    st.write(f"Barras automáticas para {col}")
                    contagem = df_filtrado[col].value_counts().reset_index()
                    contagem.columns = [col, "Contagem"]
                    st.plotly_chart(px.bar(contagem, x=col, y="Contagem"), use_container_width=True)
                if len(colunas_num) >= 2:
                    st.write("Gráfico de linha automático para duas primeiras colunas numéricas")
                    st.plotly_chart(px.line(df_filtrado[colunas_num[:2]]), use_container_width=True)
            else:
                st.info("Carregue dados e aplique filtros para gerar gráficos automáticos.")

        with aba6:
            st.subheader("Insights Automáticos")
            if not df_filtrado.empty:
                insights = gerar_insights(df_filtrado)
                st.text_area("Resumo gerado automaticamente:", insights, height=300)
            else:
                st.info("Nenhum dado para gerar insights.")

        with aba7:
            st.subheader("⬇️ Exportações (dados filtrados)")
            excel_bytes = to_excel(df_filtrado)
            csv_bytes = to_csv(df_filtrado)
            json_bytes = to_json(df_filtrado)
            st.download_button("Baixar em Excel", data=excel_bytes, file_name="dados_filtrados.xlsx", mime="application/vnd.ms-excel")
            st.download_button("Baixar em CSV", data=csv_bytes, file_name="dados_filtrados.csv", mime="text/csv")
            st.download_button("Baixar em JSON", data=json_bytes, file_name="dados_filtrados.json", mime="application/json")

        with aba8:
            st.subheader("📝 Gerar Relatório Policial (sem Anexo de Amostra)")
            with st.form("form_relatorio"):
                colA, colB = st.columns(2)
                with colA:
                    orgao = st.text_input("Órgão/Instituição", "Polícia Científica do Estado de São Paulo")
                    unidade = st.text_input("Unidade/Setor", "Núcleo de Inteligência Digital")
                    procedimento = st.text_input("Nº do Procedimento/BO", "0000000-00.0000.0.00.0000")
                    analista = st.text_input("Analista Responsável", "Perito(a) Criminal")
                with colB:
                    solicitante = st.text_input("Autoridade solicitante", "Delegado(a) de Polícia")
                    local_fuso = "America/Sao_Paulo"
                    incluir_graficos = st.checkbox("Incluir gráficos no relatório", True)
                    obs = st.text_area("Observações (opcional)")

                submitted = st.form_submit_button("Gerar Relatório")
            if submitted:
                agora = datetime.now(pytz.timezone("America/Sao_Paulo")).strftime("%d/%m/%Y %H:%M:%S %Z")
                metadados = {
                    "Órgão/Instituição": orgao,
                    "Unidade/Setor": unidade,
                    "Nº do Procedimento/BO": procedimento,
                    "Analista Responsável": analista,
                    "Autoridade Solicitante": solicitante,
                    "Local/Timezone": local_fuso,
                    "Data/Hora de Geração": agora,
                }
                if obs:
                    metadados["Observações"] = obs

                pacotes = gerar_relatorio_html_txt_docx(
                    df_base=df,
                    df_filtrado=df_filtrado if 'df_filtrado' in locals() else None,
                    incluir_graficos=incluir_graficos,
                    metadados=metadados
                )

                st.success("Relatório gerado! Baixe nos botões abaixo.")
                if pacotes.get("docx"):
                    st.download_button(
                        "Baixar Relatório (DOCX)",
                        data=pacotes["docx"],
                        file_name="relatorio_policial.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                st.download_button(
                    "Baixar Relatório (HTML)",
                    data=pacotes["html"],
                    file_name="relatorio_policial.html",
                    mime="text/html"
                )
                st.download_button(
                    "Baixar Relatório (TXT)",
                    data=pacotes["txt"],
                    file_name="relatorio_policial.txt",
                    mime="text/plain"
                )

                if not PDF_OK:
                    st.error("Para PDF, instale o pacote: pip install reportlab")
                else:
                    try:
                        pdf_bytes = gerar_relatorio_pdf(
                            df_base=df,
                            df_filtrado=df_filtrado if 'df_filtrado' in locals() else None,
                            incluir_graficos=incluir_graficos,
                            metadados=metadados,
                            titulo="Relatório Policial - Análise de IPs"
                        )
                        st.download_button(
                            "Baixar Relatório (PDF)",
                            data=pdf_bytes,
                            file_name="relatorio_policial.pdf",
                            mime="application/pdf"
                        )
                    except Exception as e:
                        st.error(f"Falha ao gerar PDF: {e}")

    else:
        st.warning("Nenhum dado válido encontrado.")


